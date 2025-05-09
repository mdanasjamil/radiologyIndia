from django.db import models
from django.conf import settings

from django.db import models
from django.conf import settings
from django.utils import timezone
from django.utils import timezone
import pytz
from django.db import models
from django.conf import settings

def ist_now():
    return timezone.now().astimezone(pytz.timezone('Asia/Kolkata'))

class Patient(models.Model):
    STATUS_CHOICES = [
        ('unassigned','Unassigned'),
        ('assigned',  'Assigned'),
        ('reported',  'Reported'),
    ]
    MODALITY_CHOICES = [
        ('MR','MR'),
        ('CT','CT'),
        ('Other','Other'),
    ]

    status             = models.CharField(max_length=20, choices=STATUS_CHOICES, default='unassigned')
    patient_id         = models.CharField(max_length=64, unique=True)
    patient_name       = models.CharField(max_length=200)
    age                = models.PositiveIntegerField(null=True, blank=True)
    sex                = models.CharField(max_length=10, choices=[('Male','Male'),('Female','Female')])
    modality           = models.CharField(max_length=10, choices=MODALITY_CHOICES)
    study_type         = models.CharField(max_length=100)
    receiving_date = models.DateTimeField(default=ist_now, editable=False)
    number_of_images   = models.PositiveIntegerField(default=0)
    preferred_physician= models.ForeignKey(
                            settings.AUTH_USER_MODEL,
                            limit_choices_to={'role':'doctor'},
                            null=True, blank=True,
                            on_delete=models.SET_NULL
                          )
    report_submitted   = models.BooleanField(default=False)
    report_file        = models.FileField(upload_to='reports/', null=True, blank=True)
    study_date         = models.DateField(null=True, blank=True)
    report_time        = models.DateTimeField(null=True, blank=True)
    institution_name   = models.CharField(max_length=200)

    def __str__(self):
        return f"{self.patient_id} – {self.patient_name}"


class DicomFile(models.Model):
    patient = models.ForeignKey(Patient, on_delete=models.CASCADE)
    file    = models.FileField(upload_to='dicom_files/')

    def __str__(self):
        return f"DICOM {self.file.name}"


class CaseFile(models.Model):
    patient = models.ForeignKey(Patient, on_delete=models.CASCADE, related_name='files')
    file = models.FileField(upload_to='cases/%Y/%m/%d')  # will store under MEDIA_ROOT/cases/YYYY/MM/DD/

    def __str__(self):
        return f"{self.patient.patient_id}: {self.file.name}"

class Report(models.Model):
    patient     = models.ForeignKey(Patient, on_delete=models.CASCADE)
    author      = models.ForeignKey(settings.AUTH_USER_MODEL, on_delete=models.CASCADE)
    report_text = models.TextField()
    created_at  = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Report by {self.author} on {self.patient_id}"

class PatientCase(models.Model):
    STATUS_CHOICES = [
        ('unassigned', 'Unassigned'),
        ('assigned', 'Assigned'),
    ]

    patient_name = models.CharField(max_length=255)
    patient_id = models.CharField(max_length=255, unique=True)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='unassigned')
    physician = models.CharField(max_length=255, blank=True, null=True)

    def __str__(self):
        return f"{self.patient_name} ({self.patient_id})"

class Report(models.Model):
    patient     = models.ForeignKey(Patient, on_delete=models.CASCADE, related_name='reports')
    author      = models.ForeignKey(settings.AUTH_USER_MODEL, on_delete=models.CASCADE)
    report_text = models.TextField()
    created_at  = models.DateTimeField(auto_now_add=True)

    class Meta:
        get_latest_by = 'created_at'

    def __str__(self):
        return f"Report #{self.id} by {self.author.username}"

    def as_doc(self):
        # returns a python-docx Document ready to save
        from docx import Document
        doc = Document()
        doc.add_heading(f'Report #{self.id}', level=1)
        doc.add_paragraph(self.report_text)
        return doc

    def filename(self):
        return f"Report_{self.id}.docx"


